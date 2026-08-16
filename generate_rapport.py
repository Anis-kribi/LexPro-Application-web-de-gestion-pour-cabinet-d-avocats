"""
Script de génération du rapport PFE LexPro amélioré en format Word (.docx)
Lit les 3 fichiers rapport_part*.md et génère Rapport_PFE_LexPro_Final.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os, re

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

def set_page_margins(doc, top=2.5, bottom=2.5, left=3.0, right=2.5):
    for section in doc.sections:
        section.top_margin = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)
        section.right_margin = Cm(right)

def add_page_number(doc):
    for section in doc.sections:
        footer = section.footer
        para = footer.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

def configure_styles(doc):
    # Normal style
    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(12)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = Pt(18)

    # Heading 1
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Times New Roman'
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x1A, 0x1A, 0x5E)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)

    # Heading 2
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Times New Roman'
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x2E, 0x4A, 0x7E)
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(8)

    # Heading 3
    h3 = doc.styles['Heading 3']
    h3.font.name = 'Times New Roman'
    h3.font.size = Pt(13)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0x35, 0x5E, 0x3E)
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(6)

def add_cover_page(doc):
    """Adds a professional cover page."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("UNIVERSITÉ — FACULTÉ DES SCIENCES")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x5E)

    doc.add_paragraph()
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("Département Informatique")
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(13)

    doc.add_paragraph()
    doc.add_paragraph()

    # PFE Title Banner
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("RAPPORT DE PROJET DE FIN D'ÉTUDES")
    run3.font.name = 'Times New Roman'
    run3.font.size = Pt(15)
    run3.font.bold = True
    run3.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)

    doc.add_paragraph()

    # Main Title
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run("LexPro")
    run4.font.name = 'Times New Roman'
    run4.font.size = Pt(32)
    run4.font.bold = True
    run4.font.color.rgb = RGBColor(0x1A, 0x1A, 0x5E)

    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run5 = p5.add_run("Application Web de Gestion de Cabinet d'Avocats")
    run5.font.name = 'Times New Roman'
    run5.font.size = Pt(16)
    run5.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph()
    doc.add_paragraph()

    # Details table
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    details = [
        ("Présenté par :", "[Votre Nom et Prénom]"),
        ("Encadrant :", "Mr. Walid"),
        ("Filière :", "Licence Informatique"),
        ("Année universitaire :", "2024 – 2025"),
    ]
    for i, (label, value) in enumerate(details):
        row = table.rows[i]
        cell_label = row.cells[0]
        cell_value = row.cells[1]
        run_l = cell_label.paragraphs[0].add_run(label)
        run_l.font.bold = True
        run_l.font.name = 'Times New Roman'
        run_l.font.size = Pt(12)
        run_v = cell_value.paragraphs[0].add_run(value)
        run_v.font.name = 'Times New Roman'
        run_v.font.size = Pt(12)

    doc.add_page_break()

def parse_and_add_content(doc, filepath):
    """Parse markdown file and add content to the docx document."""
    with open(filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # Skip horizontal rules
        if line.strip().startswith('---') and len(line.strip()) >= 3:
            i += 1
            continue

        # Detect table
        if '|' in line and i + 1 < len(lines) and '---' in lines[i+1]:
            # Collect all table rows
            table_lines = []
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i].strip())
                i += 1
            # Parse table
            rows = [r for r in table_lines if not re.match(r'^\|[-| ]+\|$', r)]
            if rows:
                num_cols = rows[0].count('|') - 1
                t = doc.add_table(rows=len(rows), cols=num_cols)
                t.style = 'Table Grid'
                for ri, row_str in enumerate(rows):
                    cells = [c.strip() for c in row_str.split('|') if c.strip() != '']
                    for ci, cell_text in enumerate(cells[:num_cols]):
                        cell = t.rows[ri].cells[ci]
                        run = cell.paragraphs[0].add_run(cell_text)
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(11)
                        if ri == 0:
                            run.font.bold = True
            continue

        # Headings
        if line.startswith('# ') and not line.startswith('## '):
            p = doc.add_heading(line[2:], level=1)
        elif line.startswith('## ') and not line.startswith('### '):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith('### ') and not line.startswith('#### '):
            p = doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            p = doc.add_heading(line[5:], level=4)

        # Blockquote (Figure placeholder)
        elif line.startswith('> '):
            content = line[2:]
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(content.replace('**', ''))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        # Bullet list
        elif line.startswith('- ') or line.startswith('* '):
            content = line[2:]
            p = doc.add_paragraph(style='List Bullet')
            # Handle **bold** in bullet
            parts = re.split(r'\*\*(.*?)\*\*', content)
            for idx, part in enumerate(parts):
                run = p.add_run(part)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                if idx % 2 == 1:
                    run.font.bold = True

        # Numbered list
        elif re.match(r'^\d+\.', line):
            content = re.sub(r'^\d+\.\s*', '', line)
            p = doc.add_paragraph(style='List Number')
            parts = re.split(r'\*\*(.*?)\*\*', content)
            for idx, part in enumerate(parts):
                run = p.add_run(part)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                if idx % 2 == 1:
                    run.font.bold = True

        # Empty line
        elif line.strip() == '':
            pass

        # Normal paragraph
        else:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(1.0)
            parts = re.split(r'\*\*(.*?)\*\*', line)
            for idx, part in enumerate(parts):
                # Also handle *italic*
                italic_parts = re.split(r'\*(.*?)\*', part)
                for jdx, ipart in enumerate(italic_parts):
                    run = p.add_run(ipart)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    if idx % 2 == 1:
                        run.font.bold = True
                    if jdx % 2 == 1:
                        run.font.italic = True

        i += 1


# ============================================================
# MAIN
# ============================================================
doc = Document()
configure_styles(doc)
set_page_margins(doc)

# Cover page
add_cover_page(doc)

# Parse all 3 parts
for part in ['rapport_part1.md', 'rapport_part2.md', 'rapport_part3.md']:
    fpath = os.path.join(BASE_DIR, part)
    if os.path.exists(fpath):
        parse_and_add_content(doc, fpath)
    else:
        print(f"Warning: {fpath} not found, skipping.")

# Page numbers
add_page_number(doc)

# ---- ADD UML DIAGRAMS SECTION ----
doc.add_page_break()
doc.add_heading("Annexe : Diagrammes UML", level=1)

uml_diagrams = [
    (
        os.path.join(BASE_DIR, 'uml_use_case.png'),
        "Diagramme de Cas d'Utilisation (Use Case)",
        "Figure A.1 : Diagramme de Cas d'Utilisation de LexPro",
        "Ce diagramme presente l'ensemble des interactions entre les acteurs du systeme (Administrateur, Avocat, Assistant) et les fonctionnalites offertes par la plateforme LexPro. Il permet d'identifier clairement le perimetre fonctionnel de l'application et les droits d'acces propres a chaque profil utilisateur."
    ),
    (
        os.path.join(BASE_DIR, 'uml_class_diagram.png'),
        "Diagramme de Classes",
        "Figure A.2 : Diagramme de Classes du Domaine LexPro",
        "Ce diagramme represente le modele de donnees de LexPro, illustrant les entites principales du systeme (User, Client, Dossier, Factures, etc.) ainsi que leurs attributs et les relations qui les unissent. Il constitue la base de la conception de la base de donnees relationnelle implementee avec l'ORM Doctrine."
    ),
    (
        os.path.join(BASE_DIR, 'uml_sequence_diagram.png'),
        "Diagramme de Sequence",
        "Figure A.3 : Diagramme de Sequence - Generation d'une Facture PDF",
        "Ce diagramme de sequence illustre le deroulement chronologique du processus de generation d'une facture au format PDF dans LexPro. Il met en evidence les interactions entre l'avocat, le navigateur web, le FacturesController, le service de verification des droits (VisibilityService) et la bibliotheque de generation PDF (Dompdf)."
    ),
]

for img_path, subtitle, figure_title, comment in uml_diagrams:
    doc.add_heading(subtitle, level=2)
    # Figure title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run(figure_title)
    run_title.font.bold = True
    run_title.font.name = 'Times New Roman'
    run_title.font.size = Pt(11)
    run_title.font.color.rgb = RGBColor(0x1A, 0x1A, 0x5E)
    # Image
    if os.path.exists(img_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_img = p_img.add_run()
        run_img.add_picture(img_path, width=Inches(6.0))
    else:
        doc.add_paragraph("[Image non trouvee: " + img_path + "]")
    # Comment
    p_comment = doc.add_paragraph()
    p_comment.paragraph_format.space_before = Pt(6)
    run_c = p_comment.add_run("Commentaire : ")
    run_c.font.bold = True
    run_c.font.name = 'Times New Roman'
    run_c.font.size = Pt(11)
    run_c2 = p_comment.add_run(comment)
    run_c2.font.name = 'Times New Roman'
    run_c2.font.size = Pt(11)
    run_c2.font.italic = True
    doc.add_paragraph()

# Save
output_path = os.path.join(BASE_DIR, 'Rapport_PFE_LexPro_v3.docx')
doc.save(output_path)
print("Rapport genere avec succes : " + output_path)
