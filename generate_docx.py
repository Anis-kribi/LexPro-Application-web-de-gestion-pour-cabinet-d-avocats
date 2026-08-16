import zlib
import base64
import urllib.request
import docx
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def get_kroki_url(puml_text):
    encoded = base64.urlsafe_b64encode(zlib.compress(puml_text.encode('utf-8'), 9)).decode('utf-8')
    return f"https://kroki.io/plantuml/png/{encoded}"

def download_image(url, filename):
    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
    with urllib.request.urlopen(req) as response, open(filename, 'wb') as out_file:
        out_file.write(response.read())

puml_class = '''@startuml
class User {
    +id: int
    +email: string
    +roles: json
}
class Client {
    +id: int
    +nom: string
    +prenom: string
    +email: string
}
class Dossier {
    +id: int
    +titre: string
    +statut: string
}
class Document {
    +id: int
    +nomDuFichier: string
}
class RendezVous {
    +id: int
    +dateRDV: date
}
class Factures {
    +id: int
    +total: float
}

Client "1" -- "*" Dossier : possède
User "1" -- "*" Dossier : gère
Dossier "1" -- "*" Document : contient
Client "1" -- "*" RendezVous : prend
Client "1" -- "*" Factures : reçoit
@enduml'''

puml_seq = '''@startuml
actor "Avocat (Admin)" as Admin
participant "Interface Web" as Vue
participant "ClientController" as Ctrl
database "Base de données" as BDD

Admin -> Vue: Clic sur "Ajouter un client"
Vue --> Admin: Affiche le formulaire d'ajout
Admin -> Vue: Saisit les informations et valide
Vue -> Ctrl: POST (données du client)
Ctrl -> Ctrl: Validation métier
Ctrl -> BDD: persist() & flush()
BDD --> Ctrl: Confirmation
Ctrl --> Vue: Redirection (succès)
Vue --> Admin: Affiche la liste actualisée
@enduml'''

print("Downloading Class Diagram...")
try:
    download_image(get_kroki_url(puml_class), "c:/symfony/Project/LexPro/classes.png")
except Exception as e:
    print(f"Failed to download class diagram: {e}")

print("Downloading Sequence Diagram...")
try:
    download_image(get_kroki_url(puml_seq), "c:/symfony/Project/LexPro/sequence.png")
except Exception as e:
    print(f"Failed to download sequence diagram: {e}")

# Start DOCX
doc = Document()
doc.styles['Normal'].font.name = 'Arial'
doc.styles['Normal'].font.size = Pt(11)

title = doc.add_heading("Rapport PFE : LexPro", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading("Sommaire", level=1)
toc = [
    "Table de figures",
    "Remerciements",
    "Introduction générale",
    "Chapitre N° 1 : Contexte Général et Spécification des Besoins",
    "Chapitre N° 2 : Conception de l'Application",
    "Chapitre N° 3 : Réalisation et Implémentation",
    "Conclusion générale"
]
for item in toc:
    doc.add_paragraph(item, style='List Bullet')
doc.add_page_break()

doc.add_heading("Table de figures", level=1)
doc.add_paragraph("Figure 1 : Diagramme de cas d'utilisation", style='List Bullet')
doc.add_paragraph("Figure 2 : Diagramme de classes", style='List Bullet')
doc.add_paragraph("Figure 3 : Diagramme de séquences (Exemple d'ajout)", style='List Bullet')
doc.add_page_break()

doc.add_heading("Remerciements", level=1)
doc.add_paragraph("Je tiens tout d'abord à remercier chaleureusement **M. Walid** pour son encadrement, ses conseils avisés et son accompagnement tout au long de ce projet. Je remercie également le jury d'avoir accepté d'évaluer ce travail, ainsi que toutes les personnes ayant contribué de près ou de loin à la réussite de ce Projet de Fin d'Études.")

# Acronymes
doc.add_heading("Liste des acronymes", level=1)
acronym_table = [
    ("PFE", "Projet de Fin d’Études"),
    ("UML", "Unified Modeling Language"),
    ("MVC", "Model‑View‑Controller"),
    ("ORM", "Object‑Relational Mapping"),
    ("SQL", "Structured Query Language"),
    ("HTML", "HyperText Markup Language"),
    ("CSS", "Cascading Style Sheets"),
    ("JS", "JavaScript"),
    ("PDF", "Portable Document Format"),
    ("API", "Application Programming Interface"),
    ("CRUD", "Create‑Read‑Update‑Delete"),
    ("XSS", "Cross‑Site Scripting"),
    ("CSRF", "Cross‑Site Request Forgery"),
    ("HTTPS", "HyperText Transfer Protocol Secure"),
    ("IDE", "Integrated Development Environment"),
    ("CI", "Continuous Integration"),
    ("GUI", "Graphical User Interface"),
    ("TTL", "Time‑to‑Live"),
    ("RBAC", "Role‑Based Access Control")
]
for abbrev, meaning in acronym_table:
    doc.add_paragraph(f"{abbrev} : {meaning}")

# Glossaire
doc.add_heading("Glossaire", level=1)
terms = [
    ("Agent juridique", "Professionnel chargé de la gestion administrative d’un cabinet (assistant·e, secrétaire)."),
    ("Facturation", "Processus d’émission de factures incluant le calcul des honoraires et frais divers."),
    ("Entrée de temps", "Enregistrement de la durée d’une activité réalisée sur un dossier, éventuellement facturable."),
    ("Dossier", "Ensemble des informations et documents liés à une affaire juridique."),
    ("Diagramme de cas d’utilisation", "UML – représente les interactions entre les acteurs externes et le système."),
    ("Diagramme de classes", "UML – modélise les entités du domaine et leurs relations."),
    ("Diagramme de séquence", "UML – décrit le déroulement chronologique d’un scénario fonctionnel.")
]
for term, definition in terms:
    doc.add_paragraph(f"{term} : {definition}")

# Bibliographie (style APA)
doc.add_heading("Bibliographie", level=1)
bibliography = [
    "Symfony Documentation. (2025). Symfony 7 – The Fast Track. Symfony. https://symfony.com/doc/7",
    "Doctrine ORM. (2024). Doctrine 3 Documentation. Doctrine Project. https://www.doctrine-project.org/projects/doctrine-orm/en/3.0",
    "Fowler, M. (2002). Patterns of Enterprise Application Architecture. Addison‑Wesley.",
    "Larman, C. (2004). Applying UML and Patterns. Prentice Hall.",
    "Clean Architecture: A Craftsman's Guide to Software Structure and Design. (2023). O’Reilly Media."
]
for entry in bibliography:
    doc.add_paragraph(entry)

# Annexes
doc.add_heading("Annexes", level=1)
annexes = [
    "Annexe A – Cahier des charges détaillé",
    "Annexe B – Maquettes d’écran (PDF)",
    "Annexe C – Extraits de code source (PHP/Symfony)",
    "Annexe D – Rapport de tests fonctionnels"
]
for a in annexes:
    doc.add_paragraph(a)

# Save the document (unchanged)
out_path = "c:/symfony/Project/LexPro/Rapport_PFE_LexPro_Final.docx"
doc.save(out_path)
print(f"Docx file saved successfully to: {out_path}")
doc.add_page_break()

doc.add_heading("Introduction générale", level=1)
doc.add_paragraph("De nos jours, l'informatique est devenue incontournable dans tous les secteurs professionnels. Le monde juridique, et particulièrement les cabinets d'avocats, n'échappent pas à cette règle. La gestion quotidienne d'un cabinet juridique nécessite une organisation minutieuse et un traitement rapide de l'information (suivi des dossiers clients, facturation, gestion des documents légaux...).")
doc.add_paragraph("C'est dans ce cadre que s'inscrit notre projet intitulé « LexPro ». Il s'agit d'une application web destinée aux professionnels du droit permettant de numériser et d'optimiser l'ensemble de leurs processus métiers. Le présent rapport détaille les différentes étapes du cycle de vie du développement de ce projet.")
doc.add_page_break()

# Chapitre 1
doc.add_heading("Chapitre N° 1 : Contexte Général et Spécification des Besoins", level=1)
doc.add_heading("Introduction", level=2)
doc.add_paragraph("Ce chapitre est consacré à la compréhension du contexte du projet et à l'analyse des besoins fonctionnels et non fonctionnels de la future application, après une analyse approfondie du système existant.")

doc.add_heading("1. Etude de l'existant", level=2)
doc.add_paragraph("Le secteur juridique est un domaine où la rigueur, l'organisation et le respect des délais sont primordiaux. Actuellement, de nombreux cabinets d'avocats gèrent encore leurs dossiers et procédures de façon manuelle ou semi-automatisée (utilisation massive de papier, traitement de texte, tableurs Excel rudimentaires). Le suivi des affaires, des rendez-vous, ainsi que la facturation dépendent fortement de l'effort administratif des assistants.")

doc.add_heading("2. Critique de l'existant (Problématique)", level=2)
doc.add_paragraph("Face à la croissance du volume de données et à l'exigence de réactivité de la part des clients, les cabinets rencontrent plusieurs difficultés :")
doc.add_paragraph("Dispersion de l'information : Les documents et notes sont souvent éparpillés, rendant la recherche d'une information spécifique longue et fastidieuse.", style='List Bullet')
doc.add_paragraph("Suivi du temps imprécis : Oublier de comptabiliser certaines heures entraîne une perte directe de chiffre d'affaires.", style='List Bullet')
doc.add_paragraph("Gestion des délais : Le non-respect d'une échéance (date d'audience) peut avoir des conséquences désastreuses.", style='List Bullet')
doc.add_paragraph("Complexité de la facturation : Générer des factures précises incluant les honoraires et les frais est souvent une tâche lourde.", style='List Bullet')

doc.add_heading("3. Solution proposée : LexPro", level=2)
doc.add_paragraph("Pour répondre à ces enjeux, nous avons conçu LexPro, une application web de gestion de cabinet d'avocats intégrée et centralisée. L'application offre une plateforme unique pour gérer les clients, les dossiers, la planification, le temps de travail et la facturation.")

doc.add_heading("4. Identification des acteurs", level=2)
doc.add_paragraph("L'application interagit principalement avec :")
doc.add_paragraph("L'Avocat / Associé : Gère ses dossiers, enregistre son temps et supervise la facturation.", style='List Bullet')
doc.add_paragraph("L'Assistant(e) Juridique : S'occupe de la saisie administrative et de la prise de rendez-vous.", style='List Bullet')
doc.add_paragraph("L'Administrateur : Gère les comptes utilisateurs.", style='List Bullet')

doc.add_heading("5. Besoins Fonctionnels Détaillés", level=2)
doc.add_paragraph("Module Tableau de bord : Afficher des indicateurs clés (dossiers en cours, tâches urgentes).", style='List Bullet')
doc.add_paragraph("Module Clients et Dossiers : Gérer le cycle de vie complet des affaires (type de cas, juridiction, adversaire).", style='List Bullet')
doc.add_paragraph("Module Gestion du Temps : Permettre aux avocats de saisir le temps passé (facturable ou non).", style='List Bullet')
doc.add_paragraph("Module Facturation : Générer automatiquement des factures basées sur les entrées de temps et gérer les articles manuels.", style='List Bullet')
doc.add_paragraph("Module Agenda/Tâches et Documents : Planifier et stocker les pièces du dossier.", style='List Bullet')

doc.add_heading("6. Besoins Non-Fonctionnels", level=2)
doc.add_paragraph("Sécurité et confidentialité (Symfony Security, protection CSRF).", style='List Bullet')
doc.add_paragraph("Ergonomie et interface responsive.", style='List Bullet')

doc.add_heading("7. Diagramme de cas d'utilisation", level=2)
doc.add_paragraph("Le diagramme suivant modélise les interactions entre les utilisateurs de LexPro et les différentes fonctionnalités :")
try:
    doc.add_picture("c:/symfony/Project/LexPro/cas_utilisation.png", width=Inches(6.0))
    last_paragraph = doc.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
except Exception as e:
    doc.add_paragraph(f"[Erreur charqement image: cas_utilisation.png]")

doc.add_page_break()

# Chapitre 2
doc.add_heading("Chapitre N° 2 : Conception de l'Application", level=1)

doc.add_heading("Introduction", level=2)
doc.add_paragraph("La conception représente l'épine dorsale du projet. Elle traduit les besoins identifiés dans le chapitre précédent en modèles informatiques clairs.")

doc.add_heading("1. Architecture Logicielle", level=2)
doc.add_paragraph("Nous avons opté pour le framework Symfony (PHP), reconnu pour sa robustesse. L'architecture repose sur le pattern MVC :")
doc.add_paragraph("Modèle (Entities Doctrine) : Gère la structure des données et la logique métier (MySQL).", style='List Bullet')
doc.add_paragraph("Vue (Twig) : Gère l'interface utilisateur de manière dynamique.", style='List Bullet')
doc.add_paragraph("Contrôleur : Fait le lien entre le Modèle et la Vue.", style='List Bullet')

doc.add_heading("2. Diagramme de classes détaillé", level=2)
doc.add_paragraph("Le modèle de données s'articule autour d'entités centrales :")
doc.add_paragraph("Dossier : Entité principale, liée à un Client et un Avocat, contenant des Tâches, Documents, EntreesDeTemps et Factures.", style='List Bullet')
doc.add_paragraph("EntreeDeTemps : Liée au Dossier, elle contient la durée et un indicateur de facturabilité.", style='List Bullet')
doc.add_paragraph("Factures et ArticleFacture : Gèrent les aspects financiers.", style='List Bullet')
try:
    doc.add_picture("c:/symfony/Project/LexPro/classes.png", width=Inches(5.0))
    last_paragraph = doc.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
except:
    pass

doc.add_heading("3. Modèle physique de données (MPD)", level=2)
doc.add_paragraph("La transformation avec l'ORM Doctrine génère les relations (One-To-Many, Many-To-One) gérées par des clés étrangères dans MySQL.")

doc.add_heading("4. Diagramme de séquences", level=2)
doc.add_paragraph("Illustration de la création d'une nouvelle entité dans le système (ex: Client) :")
try:
    doc.add_picture("c:/symfony/Project/LexPro/sequence.png", width=Inches(5.0))
    last_paragraph = doc.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
except:
    pass

doc.add_page_break()

# Chapitre 3
doc.add_heading("Chapitre N° 3 : Réalisation et Implémentation", level=1)
doc.add_heading("Introduction", level=2)
doc.add_paragraph("Ce chapitre aborde la concrétisation des modèles définis auparavant. Nous y présenterons l'environnement de développement et les interfaces de l'application.")

doc.add_heading("1. L'Environnement et Outils de développement", level=2)
doc.add_paragraph("Backend : Langage PHP 8, framework Symfony 7.", style='List Bullet')
doc.add_paragraph("SGBD : MySQL (Doctrine ORM).", style='List Bullet')
doc.add_paragraph("Frontend : Twig, CSS3, JavaScript, Bootstrap.", style='List Bullet')
doc.add_paragraph("Outils : Composer, Git, Symfony CLI.", style='List Bullet')

doc.add_heading("2. Présentation des Interfaces Utilisateur", level=2)
doc.add_paragraph("Voici un aperçu des fonctionnalités clés développées dans l'application LexPro.")

# Interface Authentification
doc.add_heading("2.1. Authentification", level=3)
doc.add_paragraph("La plateforme est sécurisée. Chaque avocat ou assistant doit se connecter avec ses identifiants.")
p = doc.add_paragraph()
r = p.add_run("[🚨 CAPTURE D'ÉCRAN À AJOUTER ICI : Page de login (login.html.twig) 🚨]")
r.font.bold = True
r.font.color.rgb = docx.shared.RGBColor(255, 0, 0)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Dashboard
doc.add_heading("2.2. Le Tableau de Bord (Dashboard)", level=3)
doc.add_paragraph("Point d'entrée de l'application, il offre une vue synthétique (statistiques, dossiers récents, tâches urgentes).")
p2 = doc.add_paragraph()
r2 = p2.add_run("[🚨 CAPTURE D'ÉCRAN À AJOUTER ICI : Tableau de bord (baseback.html.twig ou dashboard.html.twig) 🚨]")
r2.font.bold = True
r2.font.color.rgb = docx.shared.RGBColor(255, 0, 0)
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Dossiers
doc.add_heading("2.3. Gestion des Dossiers", level=3)
doc.add_paragraph("L'interface permet de consulter les détails d'une affaire (statut, priorité, tribunal). Les onglets organisent les temps, tâches et documents liés.")
p3 = doc.add_paragraph()
r3 = p3.add_run("[🚨 CAPTURE D'ÉCRAN À AJOUTER ICI : Vue détaillée d'un dossier ou liste des dossiers 🚨]")
r3.font.bold = True
r3.font.color.rgb = docx.shared.RGBColor(255, 0, 0)
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Temps
doc.add_heading("2.4. Saisie du Temps (Time Tracking)", level=3)
doc.add_paragraph("L'avocat enregistre le temps passé sur une rédaction de contrat ou une audience pour la facturation.")
p4 = doc.add_paragraph()
r4 = p4.add_run("[🚨 CAPTURE D'ÉCRAN À AJOUTER ICI : Formulaire d'ajout d'une entrée de temps 🚨]")
r4.font.bold = True
r4.font.color.rgb = docx.shared.RGBColor(255, 0, 0)
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Factures
doc.add_heading("2.5. Génération des Factures", level=3)
doc.add_paragraph("Le module automatise la création des factures en se basant sur les temps non encore facturés et les articles manuels ajoutés.")
p5 = doc.add_paragraph()
r5 = p5.add_run("[🚨 CAPTURE D'ÉCRAN À AJOUTER ICI : Formulaire de facture ou vue d'une facture générée 🚨]")
r5.font.bold = True
r5.font.color.rgb = docx.shared.RGBColor(255, 0, 0)
p5.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading("Conclusion", level=2)
doc.add_paragraph("Le résultat de cette étape est un produit fonctionnel répondant aux exigences du cabinet d'avocats.")

doc.add_page_break()

# Conclusion générale
doc.add_heading("Conclusion générale", level=1)
doc.add_paragraph("La réalisation du projet \"LexPro\" a consisté en le développement complet d'une application internet de gestion dédiée aux cabinets d'avocats. De l'analyse de l'existant à la modélisation à l'aide d'UML, jusqu'au développement PHP/Symfony final, l'objectif de fournir une solution numérisée a été atteint.")
doc.add_paragraph("Le système centralise la gestion des affaires, le suivi du temps et la facturation, apportant un gain de temps considérable pour les professionnels du droit.")
doc.add_paragraph("Ce projet de Fin d'Études (PFE) fut une expérience très enrichissante, nous permettant de maîtriser des technologies modernes telles que Symfony, Doctrine ORM et l'intégration frontend.")
doc.add_paragraph("En perspectives d'évolution futures, LexPro pourrait intégrer :")
doc.add_paragraph("Génération automatisée de documents : Pré-remplissage de PDF ou Word avec les données du dossier.", style='List Bullet')
doc.add_paragraph("Espace Client sécurisé : Pour le suivi en ligne des affaires par le client.", style='List Bullet')
doc.add_paragraph("Intégration d'API d'Agenda : Synchronisation avec Google Calendar.", style='List Bullet')

out_path = "c:/symfony/Project/LexPro/Rapport_PFE_LexPro_Final.docx"
doc.save(out_path)
print(f"Docx file saved successfully to: {out_path}")

