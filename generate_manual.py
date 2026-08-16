import docx
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
doc.styles['Normal'].font.name = 'Arial'
doc.styles['Normal'].font.size = Pt(12)

# Title
title = doc.add_heading("دليل استخدام نظام LexPro", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("هذا الدليل يشرح كيفية استخدام كل صفحة في التطبيق حسب الصلاحيات (الأدوار) المختلفة: مدير النظام (Admin)، المحامي (Avocat)، والمساعد (Assistant).\n")

# Role 1: Admin
doc.add_heading("1. مدير النظام (Administrateur)", level=1)
doc.add_paragraph("يملك المدير صلاحيات كاملة على النظام. مهمته الأساسية هي الإدارة والرقابة.")
doc.add_paragraph("• صفحة المستخدمين (Utilisateurs): إضافة محامين ومساعدين جدد، تعديل بياناتهم، وحذفهم. لا يمكن لأي شخص آخر الدخول لهذه الصفحة.", style='List Bullet')
doc.add_paragraph("• لوحة التحكم (Dashboard): رؤية إحصائيات عامة عن كل النظام.", style='List Bullet')
doc.add_paragraph("• الإعدادات (Paramètres): تعديل إعدادات المكتب (الاسم، الشعار، إلخ).", style='List Bullet')
doc.add_paragraph("• باقي الصفحات: يملك صلاحية رؤية وتعديل أي ملف، فاتورة، عميل، أو مهمة تخص أي محامي في المكتب.", style='List Bullet')

# Role 2: Avocat
doc.add_heading("2. المحامي (Avocat)", level=1)
doc.add_paragraph("المحامي هو المستخدم الأساسي للنظام، يقوم بإدارة قضاياه وعملائه.")
doc.add_paragraph("• لوحة التحكم (Dashboard): رؤية إحصائيات القضايا الخاصة به، والمهام العاجلة.", style='List Bullet')
doc.add_paragraph("• القضايا (Dossiers): إنشاء ملفات قضايا جديدة، وتتبع حالتها، وإضافة الجلسات.", style='List Bullet')
doc.add_paragraph("• العملاء (Clients): إضافة وتعديل بيانات العملاء الخاصين به.", style='List Bullet')
doc.add_paragraph("• الفواتير (Factures): إصدار فواتير للعملاء بناءً على الوقت المستغرق (Entrées de temps) والخدمات.", style='List Bullet')
doc.add_paragraph("• تتبع الوقت (Entrées de Temps): تسجيل الساعات التي قضاها في العمل على قضية معينة لتفوتر لاحقاً.", style='List Bullet')
doc.add_paragraph("• المواعيد (Rendez-vous): جدولة اجتماعات مع العملاء.", style='List Bullet')
doc.add_paragraph("• الوثائق (Documents) والمهام (Tâches): رفع المستندات وتوكيل مهام للمساعدين التابعين له.", style='List Bullet')

# Role 3: Assistant
doc.add_heading("3. المساعد (Assistant)", level=1)
doc.add_paragraph("المساعد يعمل تحت إشراف محامي معين (Manager).")
doc.add_paragraph("• لوحة التحكم (Dashboard): رؤية المهام الموكلة إليه.", style='List Bullet')
doc.add_paragraph("• العملاء (Clients): يمكنه إضافة، تعديل، وحذف العملاء نيابة عن المحامي المسؤول عنه.", style='List Bullet')
doc.add_paragraph("• الوثائق (Documents): يمكنه رفع وتعديل وحذف المستندات في القضايا التي يشرف عليها محاميه.", style='List Bullet')
doc.add_paragraph("• المهام (Tâches): يمكنه إنشاء مهام جديدة، تعديل حالتها (مكتملة، قيد الإنجاز)، وحذفها.", style='List Bullet')
doc.add_paragraph("• المواعيد (Rendez-vous): تنظيم جدول مواعيد المحامي.", style='List Bullet')
doc.add_paragraph("• الصفحات المحجوبة: لا يمكنه الوصول إلى الفواتير (Factures)، ولا يمكنه إنشاء قضايا (Dossiers) جديدة (يمكنه فقط رؤيتها)، ولا يمكنه إدارة المستخدمين.", style='List Bullet')

# Missing Features Analysis
doc.add_heading("ميزات غير موجودة (مقترحات للتطوير)", level=1)
doc.add_paragraph("بعد فحص التطبيق بشكل كامل، هذه بعض الميزات التي قد تكون ناقصة وتعتبر إضافات ممتازة في المستقبل:")

missing = [
    ("بوابة العميل (Espace Client)", "حساب خاص للعملاء للدخول ورؤية تقدم قضاياهم وتحميل مستنداتهم بأنفسهم."),
    ("عرض التقويم (Vue Calendrier)", "حالياً المواعيد تُعرض في جدول. سيكون من الأفضل عرضها في تقويم شهري/أسبوعي (مثل Google Calendar)."),
    ("تصدير الفواتير PDF (Export PDF)", "زر لتحويل الفاتورة إلى ملف PDF جاهز للطباعة وإرساله للعميل عبر الإيميل."),
    ("نظام إشعارات بالإيميل (Emails)", "إرسال إيميل تلقائي للمحامي أو العميل عند اقتراب موعد جلسة أو مهمة."),
    ("دردشة داخلية (Messagerie Interne)", "نظام محادثة بين المحامي والمساعد لتبادل الملاحظات بسرعة دون مغادرة التطبيق.")
]

for title_miss, desc in missing:
    p = doc.add_paragraph()
    r = p.add_run(f"• {title_miss}: ")
    r.font.bold = True
    r.font.color.rgb = RGBColor(220, 53, 69) # Red-ish
    p.add_run(desc)

out_path = "c:/symfony/Project/LexPro/Manuel_Utilisation_LexPro.docx"
doc.save(out_path)
print(f"Manual saved successfully to: {out_path}")
